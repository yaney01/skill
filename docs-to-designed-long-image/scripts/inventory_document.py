#!/usr/bin/env python3
"""Inventory semantic blocks from Markdown, TXT, HTML, or DOCX without rewriting copy."""

from __future__ import annotations

import argparse
import html
from html.parser import HTMLParser
import json
from pathlib import Path
import re
import sys


def normalize_space(value: str) -> str:
    return re.sub(r"\s+", " ", html.unescape(value)).strip()


def strip_inline_markdown(value: str) -> str:
    value = re.sub(r"!\[([^\]]*)\]\([^)]*\)", r"\1", value)
    value = re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", value)
    value = re.sub(r"(`+)(.*?)\1", r"\2", value)
    value = re.sub(r"(\*\*|__|~~)(.*?)\1", r"\2", value)
    value = re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)", r"\1", value)
    value = re.sub(r"(?<!_)_([^_]+)_(?!_)", r"\1", value)
    return normalize_space(value)


def parse_markdown(text: str) -> list[tuple[str, str]]:
    lines = text.replace("\r\n", "\n").replace("\r", "\n").splitlines()
    if lines and lines[0].strip() == "---":
        for index in range(1, len(lines)):
            if lines[index].strip() == "---":
                lines = lines[index + 1 :]
                break

    blocks: list[tuple[str, str]] = []
    role = "P"
    current: list[str] = []
    in_fence = False
    fence: list[str] = []

    def flush() -> None:
        nonlocal current, role
        value = strip_inline_markdown(" ".join(current))
        if value:
            blocks.append((role, value))
        current = []
        role = "P"

    for raw in lines:
        line = raw.rstrip()
        if re.match(r"^\s*```", line):
            if in_fence:
                value = normalize_space(" ".join(fence))
                if value:
                    blocks.append(("CODE", value))
                fence = []
                in_fence = False
            else:
                flush()
                in_fence = True
            continue
        if in_fence:
            fence.append(line)
            continue
        if not line.strip():
            flush()
            continue
        if re.fullmatch(r"\s{0,3}([-*_])(?:\s*\1){2,}\s*", line):
            flush()
            continue

        heading = re.match(r"^\s{0,3}(#{1,6})\s+(.+?)\s*#*\s*$", line)
        if heading:
            flush()
            blocks.append((f"H{len(heading.group(1))}", strip_inline_markdown(heading.group(2))))
            continue

        quote = re.match(r"^\s{0,3}>\s?(.*)$", line)
        if quote:
            if role != "QUOTE":
                flush()
                role = "QUOTE"
            current.append(quote.group(1))
            continue

        item = re.match(r"^\s*(?:[-+*]|\d+[.)])\s+(.+)$", line)
        if item:
            flush()
            role = "LI"
            current.append(item.group(1))
            continue

        if role == "LI" and re.match(r"^\s{2,}\S", raw):
            current.append(line.strip())
            continue
        if role in {"LI", "QUOTE"}:
            flush()
        current.append(line.strip())

    if in_fence and fence:
        blocks.append(("CODE", normalize_space(" ".join(fence))))
    flush()
    return [(role, value) for role, value in blocks if value]


class SemanticHTMLParser(HTMLParser):
    ROLE_TAGS = {
        "h1": "H1", "h2": "H2", "h3": "H3", "h4": "H4",
        "h5": "H5", "h6": "H6", "p": "P", "li": "LI",
        "blockquote": "QUOTE", "pre": "CODE", "figcaption": "CAPTION",
    }
    SKIP_TAGS = {"script", "style", "noscript", "svg"}

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.blocks: list[tuple[str, str]] = []
        self.role: str | None = None
        self.parts: list[str] = []
        self.skip_depth = 0

    def flush(self) -> None:
        value = normalize_space(" ".join(self.parts))
        if self.role and value:
            self.blocks.append((self.role, value))
        self.role = None
        self.parts = []

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        tag = tag.lower()
        if tag in self.SKIP_TAGS:
            self.skip_depth += 1
            return
        if self.skip_depth:
            return
        if tag in self.ROLE_TAGS:
            self.flush()
            self.role = self.ROLE_TAGS[tag]
        elif tag == "br" and self.role:
            self.parts.append(" ")

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()
        if tag in self.SKIP_TAGS:
            self.skip_depth = max(0, self.skip_depth - 1)
            return
        if not self.skip_depth and tag in self.ROLE_TAGS:
            self.flush()

    def handle_data(self, data: str) -> None:
        if not self.skip_depth and self.role:
            self.parts.append(data)


def parse_html(text: str) -> list[tuple[str, str]]:
    parser = SemanticHTMLParser()
    parser.feed(text)
    parser.flush()
    return parser.blocks


def parse_docx(path: Path) -> list[tuple[str, str]]:
    try:
        from docx import Document
    except ImportError as exc:
        raise SystemExit("DOCX inventory requires python-docx.") from exc

    document = Document(path)
    blocks: list[tuple[str, str]] = []
    for paragraph in document.paragraphs:
        value = normalize_space(paragraph.text)
        if not value:
            continue
        style = (paragraph.style.name or "").lower()
        heading = re.search(r"heading\s*([1-6])", style)
        if heading:
            role = f"H{heading.group(1)}"
        elif "list" in style:
            role = "LI"
        elif "quote" in style:
            role = "QUOTE"
        else:
            role = "P"
        blocks.append((role, value))
    for table in document.tables:
        for row in table.rows:
            value = " | ".join(normalize_space(cell.text) for cell in row.cells)
            if value.strip(" |"):
                blocks.append(("TABLE", value))
    return blocks


def inventory(path: Path) -> list[tuple[str, str]]:
    suffix = path.suffix.lower()
    if suffix in {".md", ".markdown", ".txt"}:
        return parse_markdown(path.read_text(encoding="utf-8"))
    if suffix in {".html", ".htm"}:
        return parse_html(path.read_text(encoding="utf-8"))
    if suffix == ".docx":
        return parse_docx(path)
    raise SystemExit(f"Unsupported input type: {suffix or '(none)'}")


def checklist_text(blocks: list[tuple[str, str]]) -> str:
    lines = [f"# 源文语义块总数：{len(blocks)}"]
    lines.extend(f"[{index:03d}][{role}] {value}" for index, (role, value) in enumerate(blocks, 1))
    return "\n".join(lines) + "\n"


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("input", nargs="?", type=Path)
    parser.add_argument("--stdin-format", choices=["markdown", "txt", "html"])
    parser.add_argument("--output", type=Path)
    parser.add_argument("--json-output", type=Path)
    args = parser.parse_args()

    if args.input:
        source = args.input.expanduser().resolve()
        if not source.is_file():
            raise SystemExit(f"Input file not found: {source}")
        blocks = inventory(source)
    elif args.stdin_format:
        raw = sys.stdin.read()
        blocks = parse_html(raw) if args.stdin_format == "html" else parse_markdown(raw)
    else:
        raise SystemExit("Provide an input file or --stdin-format.")

    if not blocks:
        raise SystemExit("No semantic blocks found.")
    rendered = checklist_text(blocks)
    if args.output:
        output = args.output.expanduser().resolve()
        output.parent.mkdir(parents=True, exist_ok=True)
        output.write_text(rendered, encoding="utf-8")
    else:
        print(rendered, end="")
    if args.json_output:
        target = args.json_output.expanduser().resolve()
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_text(
            json.dumps(
                [{"id": f"{index:03d}", "role": role, "text": value} for index, (role, value) in enumerate(blocks, 1)],
                ensure_ascii=False,
                indent=2,
            ) + "\n",
            encoding="utf-8",
        )


if __name__ == "__main__":
    main()
